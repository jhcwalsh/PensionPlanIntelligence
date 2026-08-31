"""
Text and metadata extraction from PDF and DOCX files.

Three-tier PDF strategy:
  1. pdfplumber — structured extraction (tables, layout) of the embedded
     text layer.
  2. pymupdf (fitz) — secondary text-layer extraction; tolerates a
     wider range of broken PDFs.
  3. Claude Sonnet vision — page-by-page transcription for image-only
     PDFs (scanned minutes, image-export board packs). Better at tables
     and multi-column layouts than Tesseract; costs ~$0.02–0.05 per
     multi-page document. Triggered only when both text-layer paths
     return < 100 chars of real content.
"""

import base64
import os
import re
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path

from rich.console import Console

import costs
import pdf_store
from database import (
    as_utc,
    utcnow,
    Document, ExtractionDetail, get_session, get_unextracted_documents,
)

console = Console(legacy_windows=False)

# What we KEEP and index, which has no per-token cost. This is deliberately
# not the prompt limit: summarizer.smart_truncate caps what Claude sees at
# SMART_TRUNCATE_TARGET (50k chars) regardless of how much is stored. Capping
# storage at the old 150k truncated 444 documents — 10.5% of the corpus, and
# the largest board packets — which would have capped full-text search at
# roughly their first 35 pages.
# See docs/superpowers/specs/2026-08-19-portal-readiness-design.md §2.3.
MAX_STORED_CHARS = 2_000_000

# Cap pages sent to vision OCR to bound cost on accidental 500-page agendas.
# A typical board pack is 5–60 pages; 100 covers the long tail.
MAX_VISION_OCR_PAGES = 100

# OCR-worthiness gate: vision OCR costs real money per page, so only doc
# types whose text density justifies it get the fallback — and a scan whose
# page count exceeds the doc cap is skipped outright (a 200-page image-only
# board pack isn't worth transcribing even partially).
OCR_DOC_TYPES = {"cafr", "agenda", "minutes"}

# Vision OCR is the one part of text extraction that costs money (Sonnet, per
# rendered page). Set False -- `pipeline.py --no-ocr` -- to extract everything
# a text layer will give up for free and record the rest as `ocr_deferred`,
# which `scripts/pending_spend.py` prices. Deferring loses nothing: the
# document keeps its file and its row, and re-running with OCR on picks it up.
OCR_ENABLED = True
MAX_VISION_OCR_DOC_PAGES = 50

# 2x render is roughly 1200x1600 px for letter-size — plenty of resolution for
# Sonnet vision and only modestly more image tokens than 1x.
VISION_OCR_RENDER_SCALE = 2

VISION_OCR_SYSTEM_PROMPT = (
    "You are a precision text transcriber. Output the exact text visible "
    "on the page provided, verbatim. Rules:\n"
    "- Preserve all numbers, currency symbols, percent signs, and decimal "
    "places exactly as shown.\n"
    "- Render tables using markdown pipe format, one row per line: "
    "| col1 | col2 | col3 |.\n"
    "- Preserve line breaks between paragraphs and headings.\n"
    "- Do not summarize, paraphrase, interpret, or add commentary.\n"
    "- Do not output preambles like 'Here is the text'.\n"
    "- If the page is blank or contains no readable text, output nothing."
)


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def extract_pdf_pdfplumber(path: str) -> tuple[str, int]:
    """Extract text from PDF using pdfplumber. Returns (text, page_count)."""
    try:
        import pdfplumber
        pages_text = []
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                # Also extract tables as plain text
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = " | ".join(str(cell or "").strip() for cell in row)
                        if row_text.strip():
                            text += "\n" + row_text
                if text.strip():
                    pages_text.append(f"[Page {i + 1}]\n{text}")
        full_text = "\n\n".join(pages_text)
        return full_text[:MAX_STORED_CHARS], page_count
    except Exception as e:
        console.print(f"  [yellow]pdfplumber failed: {e}, trying pymupdf...[/yellow]")
        return "", 0


def extract_pdf_pymupdf(path: str) -> tuple[str, int]:
    """Fallback PDF extraction using pymupdf (fitz)."""
    try:
        import fitz  # pymupdf
        doc = fitz.open(path)
        page_count = len(doc)
        pages_text = []
        for i, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages_text.append(f"[Page {i + 1}]\n{text}")
        full_text = "\n\n".join(pages_text)
        return full_text[:MAX_STORED_CHARS], page_count
    except Exception as e:
        console.print(f"  [red]pymupdf also failed: {e}[/red]")
        return "", 0


@dataclass
class OcrInfo:
    """How far vision OCR actually got (for the extraction_details index)."""
    pages_ocred: int = 0
    # 'page_cap'    the whole-doc gate fired
    # 'api_error'   every page failed at the API, so nothing was learned about
    #               the document. Distinct because the alternative is recording
    #               'ocr_empty', which asserts the document has no text.
    reason: str | None = None


def extract_pdf_ocr(path: str) -> tuple[str, int, OcrInfo]:
    """OCR fallback using Claude Sonnet vision.

    Renders each page with pymupdf and asks Sonnet for verbatim
    transcription. Replaces the prior Tesseract path: better at tables,
    multi-column layouts, and scanned forms, at a cost of ~$0.02–0.05
    per multi-page document. Failure on a single page (network blip,
    transient API error) is logged and skipped — other pages still
    contribute. The function returns whatever was successfully
    transcribed; an empty result causes ``extract_document`` to mark
    the row ``failed`` as before.

    Every Claude call inside is labelled ``ocr``: vision is billed per page as
    image input, which makes this the leading suspect for the largest line
    item on the bill. Labelling it is what lets the Spend tab confirm or clear
    that, rather than leaving it a guess. Applied to the whole function rather
    than the one call, so the per-page loop needs no re-indenting and a second
    call added later is labelled automatically.
    """
    with costs.track("ocr"):
        return _extract_pdf_ocr(path)


def _extract_pdf_ocr(path: str) -> tuple[str, int, OcrInfo]:
    try:
        import fitz  # pymupdf
    except ImportError as e:
        console.print(f"  [yellow]OCR skipped: pymupdf not installed ({e})[/yellow]")
        return "", 0, OcrInfo()

    try:
        doc = fitz.open(path)
    except Exception as e:
        console.print(f"  [red]Vision OCR failed: {e}[/red]")
        return "", 0, OcrInfo()

    if len(doc) > MAX_VISION_OCR_DOC_PAGES:
        console.print(
            f"  [yellow]OCR skipped: {len(doc)} pages exceeds the "
            f"{MAX_VISION_OCR_DOC_PAGES}-page document cap[/yellow]"
        )
        return "", len(doc), OcrInfo(reason="page_cap")

    try:
        from summarizer import MODEL_SONNET, _get_client
    except ImportError as e:
        console.print(f"  [yellow]OCR skipped: anthropic SDK not available ({e})[/yellow]")
        return "", 0, OcrInfo()

    try:
        client = _get_client()
    except Exception as e:
        console.print(f"  [yellow]OCR skipped: no Anthropic credentials ({e})[/yellow]")
        return "", 0, OcrInfo()

    try:
        page_count = len(doc)
        pages_text = []
        pages_attempted = 0
        api_errors = 0
        mat = fitz.Matrix(VISION_OCR_RENDER_SCALE, VISION_OCR_RENDER_SCALE)
        for i, page in enumerate(doc):
            if i >= MAX_VISION_OCR_PAGES:
                console.print(
                    f"  [yellow]Vision OCR cap reached at page {MAX_VISION_OCR_PAGES}; "
                    f"skipping remaining {page_count - i} pages[/yellow]"
                )
                break
            pages_attempted = i + 1
            pix = page.get_pixmap(matrix=mat)
            png_b64 = base64.b64encode(pix.tobytes("png")).decode("ascii")
            try:
                msg = client.messages.create(
                    model=MODEL_SONNET,
                    max_tokens=4096,
                    system=VISION_OCR_SYSTEM_PROMPT,
                    messages=[{
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/png",
                                    "data": png_b64,
                                },
                            },
                            {"type": "text", "text": f"Transcribe page {i + 1}."},
                        ],
                    }],
                )
                page_text = msg.content[0].text if msg.content else ""
            except Exception as e:
                console.print(f"  [red]Vision OCR failed on page {i + 1}: {e}[/red]")
                api_errors += 1
                continue
            if page_text.strip():
                pages_text.append(f"[Page {i + 1}]\n{page_text}")
        full_text = "\n\n".join(pages_text)
        # Every page failed at the API, so this run learned nothing about the
        # document. Saying so is the difference between "retry when the key
        # works" and "this document has no text in it".
        if api_errors and api_errors == pages_attempted:
            return "", page_count, OcrInfo(pages_ocred=0, reason="api_error")
        return full_text[:MAX_STORED_CHARS], page_count, OcrInfo(pages_ocred=pages_attempted)
    except Exception as e:
        console.print(f"  [red]Vision OCR failed: {e}[/red]")
        return "", 0, OcrInfo(reason="api_error")


def _looks_like_pdf(path: str) -> bool:
    """True if the bytes start with the PDF magic number.

    The extension is what the URL claimed; this is what arrived. A plan that
    serves an error page, a login wall or a JS redirect to a .pdf URL leaves a
    file whose name says PDF and whose first five bytes say '<!DOC'.
    """
    try:
        with open(path, "rb") as f:
            return f.read(5) == b"%PDF-"
    except OSError:
        return False


def extract_pdf(path: str, allow_ocr: bool = True,
                gate_reason: str = "ocr_gate_doc_type"
                ) -> tuple[str, int, str | None, int | None]:
    """Extract PDF text. Returns (text, pages, reason, pages_ocred) where
    reason is an extraction_details reason for empty/partial results.

    `gate_reason` names *which* gate stopped OCR, so a document deferred to
    keep spend down is distinguishable from one whose doc_type was never
    OCR-worthy. Both leave the document re-processable; only one is a
    decision anyone would want to revisit.
    """
    # Check what the file actually is before deciding what it failed to say.
    # Thirteen documents in the OCR backlog turned out to be HTML saved with a
    # .pdf extension -- an error or redirect page the fetcher stored without
    # noticing. Every one was recorded 'ocr_empty', which asserts OCR read a
    # scanned document and found nothing in it. They were never PDFs, so no
    # amount of OCR was ever going to help, and the label hid a fetcher
    # problem behind an extraction one.
    if not _looks_like_pdf(path):
        return "", 0, "not_a_pdf", None

    text, pages = extract_pdf_pdfplumber(path)
    if len(text.strip()) < 100:
        text, pages = extract_pdf_pymupdf(path)
    if len(text.strip()) >= 100:
        return text, pages, None, None
    if not allow_ocr:
        return text, pages, gate_reason, None
    console.print("  [dim]Trying OCR...[/dim]")
    text, pages, info = extract_pdf_ocr(path)
    if not text.strip():
        if info.reason == "page_cap":
            reason = "ocr_gate_page_cap"
        elif info.reason == "api_error":
            # Not a fact about the document. On 2026-08-31 an exhausted
            # Anthropic credit balance turned 30 ocr_deferred rows into
            # ocr_empty, dropping them out of the priced backlog that
            # scripts/pending_spend.py reports — the work became invisible
            # rather than pending.
            reason = "ocr_unavailable"
        else:
            reason = "ocr_empty"
        return text, pages, reason, info.pages_ocred
    if info.pages_ocred < pages:
        return text, pages, "ocr_partial", info.pages_ocred
    return text, pages, None, None


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def extract_docx(path: str) -> tuple[str, int]:
    """Extract text from a Word document."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        full_text = "\n".join(paragraphs)
        return full_text[:MAX_STORED_CHARS], 1  # DOCX doesn't have "pages" in the same way
    except Exception as e:
        console.print(f"  [red]DOCX extraction failed: {e}[/red]")
        return "", 0


# ---------------------------------------------------------------------------
# Metadata helpers
# ---------------------------------------------------------------------------

MEETING_TYPE_PATTERNS = {
    "investment": [r"investment\s+committee", r"investment\s+board", r"portfolio"],
    "audit": [r"audit\s+committee", r"risk\s+committee"],
    "board": [r"board\s+of\s+(trustees|directors|retirement)", r"full\s+board"],
    "actuarial": [r"actuarial", r"funded\s+status"],
}


def infer_meeting_type(text: str, filename: str) -> str:
    combined = (text[:2000] + " " + filename).lower()
    for mtype, patterns in MEETING_TYPE_PATTERNS.items():
        for p in patterns:
            if re.search(p, combined, re.IGNORECASE):
                return mtype
    return "board"


DATE_RE = re.compile(
    r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)"
    r"\s+\d{1,2},?\s+\d{4}\b",
    re.IGNORECASE
)


_FILENAME_DATE_PATTERNS = [
    # 04292026 / 04-29-2026 / 04_29_2026 / 04.29.2026 — concatenated MMDDYYYY
    # with optional separators. Matches "agenda.board.04292026.pdf",
    # "Board_Pack_04-29-2026.pdf", etc.
    (r"(?:^|[^0-9])(\d{2})[\-_.]?(\d{2})[\-_.]?(\d{4})(?:[^0-9]|$)", "MDY"),
    # M.D.YY or M-D-YY (two-digit year). Matches "IC_Agenda_4.24.26.pdf".
    # Constrained to . / - separators so we don't snag random number runs.
    (r"(?:^|[^0-9])(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{2})(?:[^0-9]|$)", "MDY2"),
    # YYYYMMDD / YYYY-MM-DD / YYYY_MM_DD
    (r"(?:^|[^0-9])(\d{4})[\-_.]?(\d{2})[\-_.]?(\d{2})(?:[^0-9]|$)", "YMD"),
    # Month DD YYYY (with various separators)
    (r"(January|February|March|April|May|June|July|August|September|October|"
     r"November|December)[\s\-_]+(\d{1,2})[\s\-_,]+(\d{4})", "WORD_MDY"),
    # Month YYYY only (no day) — fall back to first of month.
    # Matches "December-2025-Board-Highlights.pdf".
    (r"(January|February|March|April|May|June|July|August|September|October|"
     r"November|December)[\s\-_]+(\d{4})", "WORD_MY"),
]


def parse_date_from_filename(filename: str | None) -> datetime | None:
    """Extract a plausible meeting date from a filename.

    Tries several common shapes (MMDDYYYY without separators, M.D.YY,
    YYYYMMDD, "Month DD YYYY", "Month YYYY") and returns the first
    valid result. Returns None if no plausible date is found.

    The fetcher's ``parse_date_from_text`` only handles separator-bearing
    formats, which misses board-agenda filenames like
    ``agenda.board.04292026.pdf``. This helper is the second-pass fallback.
    """
    if not filename:
        return None
    base = re.sub(r"\.[a-zA-Z0-9]+$", "", filename)
    for pattern, kind in _FILENAME_DATE_PATTERNS:
        m = re.search(pattern, base, re.IGNORECASE)
        if not m:
            continue
        try:
            if kind == "MDY":
                month, day, year = int(m.group(1)), int(m.group(2)), int(m.group(3))
            elif kind == "MDY2":
                month, day, year = int(m.group(1)), int(m.group(2)), int(m.group(3))
                year = year + 2000 if year < 100 else year
            elif kind == "YMD":
                year, month, day = int(m.group(1)), int(m.group(2)), int(m.group(3))
            elif kind == "WORD_MDY":
                return datetime.strptime(
                    f"{m.group(1)} {m.group(2)} {m.group(3)}", "%B %d %Y"
                )
            elif kind == "WORD_MY":
                return datetime.strptime(f"{m.group(1)} 1 {m.group(2)}", "%B %d %Y")
            else:
                continue
            if (1 <= month <= 12 and 1 <= day <= 31
                    and 2000 <= year <= 2035):
                return datetime(year, month, day)
        except ValueError:
            continue
    return None


def _date_is_plausible(d: datetime, downloaded_at: datetime | None) -> bool:
    """Reject dates that can't possibly be a real meeting given fetch time.

    A meeting date should be within ~60 days after download (forward-scheduled
    agendas) and within ~5 years before (the longest historical material we
    routinely fetch). Anything outside that window is almost certainly a
    parser misread.
    """
    if downloaded_at is None:
        return True
    # `d` is parsed out of text or a filename and is always naive;
    # downloaded_at comes from the DB and is aware on Postgres. Normalise both
    # rather than assume either shape.
    d = as_utc(d)
    downloaded_at = as_utc(downloaded_at)
    if d > downloaded_at + timedelta(days=60):
        return False
    if d < downloaded_at - timedelta(days=5 * 365):
        return False
    return True


def infer_meeting_date(
    text: str,
    existing_date: datetime | None,
    filename: str | None = None,
    downloaded_at: datetime | None = None,
) -> datetime | None:
    """Best-effort meeting date for a Document.

    Priority order:
      1. ``existing_date`` if already set and plausible (fetcher wins)
      2. ``parse_date_from_filename`` — strong signal, low false-positive rate
      3. First Month-DD-YYYY in the first 2000 chars of ``text`` — last resort

    Sanity-checks every candidate against ``downloaded_at`` (when provided):
    a date >60 days after fetch or >5 years before is treated as a parser
    error and discarded. Returns None when no plausible date is found —
    "no date" is preferred over a wrong date.
    """
    if existing_date and _date_is_plausible(existing_date, downloaded_at):
        return existing_date

    fname_date = parse_date_from_filename(filename)
    if fname_date and _date_is_plausible(fname_date, downloaded_at):
        return fname_date

    m = DATE_RE.search(text[:2000])
    if m:
        raw = m.group(0).replace(",", "").strip()
        for fmt in ("%B %d %Y", "%B %d, %Y"):
            try:
                d = datetime.strptime(raw, fmt)
                if _date_is_plausible(d, downloaded_at):
                    return d
            except ValueError:
                continue
    return None


# ---------------------------------------------------------------------------
# Main extraction runner
# ---------------------------------------------------------------------------

@dataclass
class ExtractOutcome:
    """Result of one document extraction, including why it fell short.

    ``reason`` is non-None whenever the document is not fully extracted —
    see database.ExtractionDetail for the vocabulary. A ``done`` outcome
    can still carry reason='ocr_partial'.
    """
    text: str = ""
    pages: int = 0
    status: str = "failed"
    reason: str | None = None
    pages_ocred: int | None = None


def extract_document(doc: Document) -> ExtractOutcome:
    """Extract text from a document's PDF, wherever it lives.

    Reads through ``pdf_store.document_pdf``, so a document whose local file
    is gone but whose bytes were retained extracts normally. Before this,
    extraction only ever worked on the machine that did the fetching — the
    coupling CLAUDE.md records as having caused two separate defects, and the
    reason 128 documents sit at `file_missing` with nothing to re-extract.

    ``file_missing`` now means genuinely nowhere: not on disk *and* not in
    the store.
    """
    try:
        with pdf_store.document_pdf(doc) as pdf_path:
            return _extract_from_path(doc, str(pdf_path))
    except FileNotFoundError:
        return ExtractOutcome(reason="file_missing")


def _extract_from_path(doc: Document, path: str) -> ExtractOutcome:
    """Extract from a readable path, whether local or pulled from the store.

    Takes the path rather than reading ``doc.local_path`` so the retained
    copy and the local file follow exactly the same code.
    """
    # From the document, never from `path`. A copy pulled out of the store is
    # always written to a `.pdf` temp file regardless of what it holds, so
    # reading the suffix off the resolved path would hand every retained
    # .docx to the PDF extractor.
    name = doc.filename or Path(doc.local_path or path).name
    ext = Path(name).suffix.lower()

    console.print(f"  Extracting [cyan]{name}[/cyan]")

    if ext == ".pdf":
        type_allows_ocr = doc.doc_type in OCR_DOC_TYPES
        text, pages, reason, pages_ocred = extract_pdf(
            path,
            allow_ocr=OCR_ENABLED and type_allows_ocr,
            # Only call it deferred when spend is the reason: a doc_type that
            # was never OCR-worthy is not work anyone is waiting to fund.
            gate_reason=("ocr_deferred" if type_allows_ocr and not OCR_ENABLED
                         else "ocr_gate_doc_type"),
        )
    elif ext in (".docx", ".doc"):
        text, pages = extract_docx(path)
        reason, pages_ocred = ("extract_empty" if not text.strip() else None), None
    else:
        return ExtractOutcome(reason="unsupported_format")

    if not text.strip():
        return ExtractOutcome(pages=pages, reason=reason or "extract_empty",
                              pages_ocred=pages_ocred)

    return ExtractOutcome(text=text, pages=pages, status="done",
                          reason=reason, pages_ocred=pages_ocred)


def _apply_outcome(session, doc, outcome):
    """Write one extraction result onto `doc`. Idempotent: safe to re-run
    against a freshly-loaded copy after a rollback."""
    # On failure, leave extracted_text alone: never store "" (the
    # GzippedText wrapper would persist it as a non-NULL gzip blob),
    # and never clobber text kept from an earlier successful pass.
    if outcome.status == "done":
        doc.extracted_text = outcome.text
        doc.page_count = outcome.pages
    doc.extraction_status = outcome.status

    # Keep the extraction_details index in sync: any shortfall (failure or
    # partial scan) is recorded so the doc can be found and re-processed
    # later; a clean pass clears it.
    if outcome.reason:
        session.merge(ExtractionDetail(
            document_id=doc.id, reason=outcome.reason,
            pages_total=outcome.pages or None,
            pages_ocred=outcome.pages_ocred,
            detected_at=utcnow()))
    else:
        session.query(ExtractionDetail).filter(
            ExtractionDetail.document_id == doc.id).delete()

    # Try to infer meeting date from content if not already set
    if outcome.text:
        doc.meeting_date = infer_meeting_date(
            outcome.text, doc.meeting_date,
            filename=doc.filename,
            downloaded_at=doc.downloaded_at,
        )


def _persist_outcome(session, doc, outcome):
    """Commit one extraction result, reconnecting once if the link died.

    Extraction happens between commits, so the connection sits idle for as
    long as the parse takes -- minutes on a large board pack. Neon closes it
    in that window and the write then lands on a dead socket:

        psycopg.OperationalError: SSL connection has been closed unexpectedly

    Found on a 2,596-page, 20.5 MB MCERA packet, which failed reproducibly
    and took the whole plan's run down with it (the loop is unguarded, so one
    document's failure stranded the nine queued behind it). The write itself
    is not the problem -- the same payload commits in 0.1s on a live
    connection -- so a rollback, a fresh load and one retry clears it.

    The same idle window surfaces two different exceptions depending on who
    notices first, and both have now been seen live on this corpus:

        OperationalError  psycopg.OperationalError:
                          SSL connection has been closed unexpectedly
        InternalError     psycopg.errors.IdleInTransactionSessionTimeout:
                          terminating connection due to idle-in-transaction
                          timeout

    Catching only the first is what a plausible-looking fix does; it then
    fails on the next large document with an unfamiliar traceback. Both are
    caught here.

    `pool_pre_ping` does not help either: it validates at checkout, and this
    connection is checked out and healthy before the parse begins.
    """
    from sqlalchemy.exc import InternalError, OperationalError

    try:
        _apply_outcome(session, doc, outcome)
        session.commit()
        return doc
    except (OperationalError, InternalError) as e:
        console.print(f"  [yellow]connection lost during extraction, "
                      f"retrying commit: {str(e)[:90]}[/yellow]")
        session.rollback()
        doc = session.get(Document, doc.id)
        _apply_outcome(session, doc, outcome)
        session.commit()
        return doc


def run_extractor(doc_ids: list[int] = None, retry_failed: bool = False):
    """
    Extract text for all pending documents (or specific doc_ids).
    Updates extraction_status, extracted_text, page_count in DB.
    Pass retry_failed=True to re-attempt previously failed documents.
    """
    session = get_session()
    try:
        if doc_ids:
            docs = session.query(Document).filter(Document.id.in_(doc_ids)).all()
        elif retry_failed:
            docs = session.query(Document).filter(
                Document.extraction_status == "failed"
            ).all()
            # Reset to pending so the run loop can update them
            for doc in docs:
                doc.extraction_status = "pending"
            session.commit()
        else:
            docs = get_unextracted_documents(session)

        if not docs:
            console.print("[yellow]No documents pending extraction.[/yellow]")
            return

        console.print(f"[bold]Extracting text from {len(docs)} documents...[/bold]")

        for doc in docs:
            outcome = extract_document(doc)
            doc = _persist_outcome(session, doc, outcome)
            status_color = "green" if outcome.status == "done" else "red"
            note = f" [{outcome.reason}]" if outcome.reason else ""
            console.print(f"    [{status_color}]{outcome.status}[/{status_color}] "
                          f"— {outcome.pages} pages, {len(outcome.text):,} chars{note}")

        done = sum(1 for d in docs if d.extraction_status == "done")
        console.print(f"\n[bold green]{done}/{len(docs)} documents extracted successfully.[/bold green]")

    finally:
        session.close()


if __name__ == "__main__":
    run_extractor()
